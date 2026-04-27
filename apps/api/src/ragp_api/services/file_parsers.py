"""Parsers that convert raw file bytes to plain text for indexing."""

from __future__ import annotations

import io
import os

from fastapi import HTTPException


def parse_to_text(filename: str, content_type: str, raw: bytes) -> str:
    """Return extracted plain text from *raw* bytes.

    Dispatches to a specialised parser based on file extension / MIME type.
    Falls back to UTF-8 decode for plain-text formats.

    Raises:
        HTTPException(422): extraction failed or yielded no usable text.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf" or content_type == "application/pdf":
        return _parse_pdf(raw)

    if ext == ".docx" or content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _parse_docx(raw)

    # Plain-text fallback
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="File must be valid UTF-8") from exc


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _parse_pdf(raw: bytes) -> str:
    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover
        raise HTTPException(
            status_code=500, detail="pypdf is not installed; cannot parse PDF"
        ) from exc

    reader = pypdf.PdfReader(io.BytesIO(raw))
    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if page_text:
            pages.append(page_text)

    if not pages:
        raise HTTPException(
            status_code=422,
            detail="PDF has no extractable text (likely a scanned image)",
        )

    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _parse_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise HTTPException(
            status_code=500, detail="python-docx is not installed; cannot parse DOCX"
        ) from exc

    doc = Document(io.BytesIO(raw))

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    if not parts:
        raise HTTPException(
            status_code=422,
            detail="DOCX has no extractable text",
        )

    return "\n\n".join(parts)
